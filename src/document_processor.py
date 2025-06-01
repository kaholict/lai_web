import os
import logging
from typing import List, Dict
from pathlib import Path
import PyPDF2
import docx
import docxpy
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Класс для обработки PDF и DOCX документов"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Извлечение текста из PDF файла"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            logger.info(f"Успешно извлечен текст из PDF: {file_path}")
            return text
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из PDF {file_path}: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Извлечение текста из DOCX файла с поддержкой гиперссылок"""
        try:
            # Сначала пытаемся использовать docxpy для извлечения гиперссылок
            text = ""
            hyperlinks = []
            
            try:
                # Используем docxpy для извлечения текста и гиперссылок
                doc = docxpy.DOCReader(file_path)
                doc.process()
                
                text = docxpy.process(file_path)
                
                # Безопасная обработка гиперссылок
                if hasattr(doc, 'data') and doc.data:
                    raw_links = doc.data.get('links', [])
                    
                    # Обрабатываем различные форматы данных о гиперссылках
                    for link in raw_links:
                        try:
                            if isinstance(link, dict):
                                # Если гиперссылка - словарь
                                link_text = link.get('text', '')
                                link_url = link.get('url', '')
                            elif isinstance(link, (tuple, list)) and len(link) >= 2:
                                # Если гиперссылка - кортеж или список
                                link_text = str(link[0]) if len(link) > 0 else ''
                                link_url = str(link[1]) if len(link) > 1 else ''
                            elif hasattr(link, 'text') and hasattr(link, 'url'):
                                # Если гиперссылка - объект с атрибутами
                                link_text = str(link.text) if link.text else ''
                                link_url = str(link.url) if link.url else ''
                            else:
                                # Пытаемся преобразовать в строку
                                link_text = str(link)
                                link_url = ""
                            
                            if link_text or link_url:
                                hyperlinks.append((link_text, link_url))
                        except Exception as link_error:
                            logger.warning(f"Ошибка при обработке гиперссылки: {link_error}")
                            continue
                
            except Exception as docxpy_error:
                logger.warning(f"Ошибка при использовании docxpy для {file_path}: {docxpy_error}")
                # Fallback к стандартной библиотеке для извлечения текста
                text = self._extract_text_with_python_docx(file_path)
            
            # Если текст не извлечен, используем fallback
            if not text or not text.strip():
                text = self._extract_text_with_python_docx(file_path)
            
            # Добавляем информацию о гиперссылках к тексту
            if hyperlinks:
                text += "\n\nГиперссылки в документе:\n"
                for link_text, link_url in hyperlinks:
                    if link_text and link_url:
                        text += f"- {link_text}: {link_url}\n"
                    elif link_url:
                        text += f"- {link_url}\n"
                    elif link_text:
                        text += f"- {link_text}\n"
            
            logger.info(f"Успешно извлечен текст из DOCX: {file_path}")
            return text
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении текста из DOCX {file_path}: {e}")
            # Окончательный fallback к стандартной библиотеке
            return self._extract_text_with_python_docx(file_path)
    
    def _extract_text_with_python_docx(self, file_path: str) -> str:
        """Извлечение текста с помощью стандартной библиотеки python-docx"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Пытаемся извлечь гиперссылки с помощью python-docx
            try:
                hyperlinks = self._extract_hyperlinks_python_docx(doc)
                if hyperlinks:
                    text += "\n\nГиперссылки в документе:\n"
                    for link_text, link_url in hyperlinks:
                        if link_text and link_url:
                            text += f"- {link_text}: {link_url}\n"
                        elif link_url:
                            text += f"- {link_url}\n"
            except Exception as hyperlink_error:
                logger.warning(f"Не удалось извлечь гиперссылки: {hyperlink_error}")
            
            return text
            
        except Exception as e:
            logger.error(f"Fallback извлечение текста также не удалось для {file_path}: {e}")
            return ""
    
    def _extract_hyperlinks_python_docx(self, doc) -> List[tuple]:
        """Извлечение гиперссылок с помощью python-docx"""
        hyperlinks = []
        
        try:
            # Извлекаем гиперссылки из отношений документа
            rels = doc.part.rels
            for rel in rels:
                if rels[rel].reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink":
                    url = rels[rel]._target
                    
                    # Ищем текст гиперссылки в документе
                    link_text = ""
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            if hasattr(run.element, 'hyperlink') and run.element.hyperlink:
                                link_text = run.text
                                break
                        if link_text:
                            break
                    
                    hyperlinks.append((link_text, url))
        except Exception as e:
            logger.warning(f"Ошибка при извлечении гиперссылок через python-docx: {e}")
        
        return hyperlinks
    
    def process_documents(self, folder_path: str) -> List[Document]:
        """Обработка всех документов в папке"""
        documents = []
        folder = Path(folder_path)
        
        if not folder.exists():
            logger.error(f"Папка не существует: {folder_path}")
            return documents
        
        processed_files = 0
        total_files = 0
        
        for file_path in folder.rglob("*"):
            if file_path.is_file():
                total_files += 1
                text = ""
                metadata = {
                    "source": str(file_path),
                    "filename": file_path.name,
                    "file_type": file_path.suffix.lower()
                }
                
                if file_path.suffix.lower() == ".pdf":
                    text = self.extract_text_from_pdf(str(file_path))
                elif file_path.suffix.lower() == ".docx":
                    text = self.extract_text_from_docx(str(file_path))
                
                if text and text.strip():
                    # Разбиваем текст на чанки
                    chunks = self.text_splitter.split_text(text)
                    for i, chunk in enumerate(chunks):
                        chunk_metadata = metadata.copy()
                        chunk_metadata["chunk_id"] = i
                        documents.append(Document(
                            page_content=chunk,
                            metadata=chunk_metadata
                        ))
                    processed_files += 1
                else:
                    logger.warning(f"Не удалось извлечь текст из файла: {file_path}")
        
        logger.info(f"Обработано {len(documents)} чанков из {processed_files}/{total_files} файлов")
        return documents
